# -*- coding: utf-8 -*-
from pathlib import Path
import re
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
JOURNAL_DIR = ROOT / "업무일지"
OUTPUT = Path(__file__).resolve().parent / "완속충전기_통합업무기록_2026-07-06_현재.docx"

NAVY = RGBColor(21, 63, 74)
TEAL = RGBColor(20, 107, 93)
MUTED = RGBColor(102, 116, 124)
LIGHT = "EAF3F1"
GRAY = "F2F4F7"
WHITE = RGBColor(255, 255, 255)


def set_font(run, size=None, bold=None, color=None, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_inches) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=9, color=MUTED)


def clean_inline(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def read_docs():
    docs = []
    for path in sorted(JOURNAL_DIR.glob("2026-*.md")):
        if "종합업무일지" in path.name:
            continue
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="cp949", errors="replace")
        docs.append((path.stem, text))
    return docs


def parse_calendar(text):
    rows = []
    for line in text.splitlines():
        if re.match(r"^\|\s*\d+\s*\|", line):
            cells = [clean_inline(c) for c in line.split("|")[1:-1]]
            if len(cells) >= 4:
                rows.append(cells[:4])
    return rows


def parse_markdown(text):
    blocks = []
    table = []
    for raw in text.replace("\r", "").split("\n"):
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            table.append(line)
            continue
        if table:
            blocks.append(("table", table))
            table = []
        if not line.strip():
            continue
        if line.startswith("### "):
            blocks.append(("h3", clean_inline(line[4:])))
        elif line.startswith("## "):
            blocks.append(("h2", clean_inline(line[3:])))
        elif line.startswith("# "):
            blocks.append(("h1", clean_inline(line[2:])))
        elif line.startswith("- "):
            blocks.append(("bullet", clean_inline(line[2:])))
        else:
            blocks.append(("p", clean_inline(line)))
    if table:
        blocks.append(("table", table))
    return blocks


def add_table_from_markdown(doc, lines):
    rows = []
    for idx, line in enumerate(lines):
        if idx == 1 and re.fullmatch(r"\|[\s\-:|]+\|", line):
            continue
        rows.append([clean_inline(c) for c in line.split("|")[1:-1]])
    if not rows or len(rows[0]) < 2:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    if cols == 4:
        widths = [0.55, 1.65, 1.05, 3.25]
    elif cols == 3:
        widths = [1.3, 2.0, 3.2]
    elif cols == 2:
        widths = [1.55, 4.95]
    else:
        widths = [6.5 / cols] * cols
    set_table_widths(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=9.2, bold=(r_idx == 0), color=(NAVY if r_idx == 0 else None))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
    set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build():
    docs = read_docs()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "완속충전기 통합 업무기록"
    set_font(header.runs[0], size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.add_run("2026-07-06 ~ 현재   |   ")
    set_font(footer.runs[0], size=9, color=MUTED)
    add_page_number(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(24)
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("PROJECT WORK RECORD")
    set_font(run, size=10, bold=True, color=TEAL)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("완속충전기 통합 업무기록")
    set_font(run, size=28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("ChatGPT Work + Codex 대화 기반 · 날짜별 업무 목적·진행·결과 통합본")
    set_font(run, size=12.5, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    set_table_widths(meta, [1.45, 5.05])
    metadata = [
        ("정리 기간", "2026-07-06 ~ 2026-07-24"),
        ("통합 원천", "ChatGPT Work 대화, Codex 대화, 프로젝트 업무일지"),
        ("작성 원칙", "코딩 세부사항 제외 · 업무 목적/달성 방법/결과 중심 · 확인되지 않은 사실은 확인 필요로 표시"),
    ]
    for idx, (label, value) in enumerate(metadata):
        for col, text in enumerate((label, value)):
            cell = meta.cell(idx, col)
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            set_font(run, size=9.5, bold=(col == 0), color=(NAVY if col == 0 else None))
            if col == 0:
                set_cell_shading(cell, LIGHT)

    doc.add_paragraph()
    doc.add_heading("통합 요약", level=1)
    total_conversations = sum(len(re.findall(r"^###\s+대화", text, flags=re.MULTILINE)) for _, text in docs)
    summary_p = doc.add_paragraph()
    summary_p.add_run(
        f"총 {len(docs)}개 날짜의 기록과 {total_conversations}개 대화 단위를 하나로 통합했다. "
        "충전소 후보·목록·측정보고서의 정합성 점검, MapWebApp 운영 개선, 현장 계측 관리, "
        "이재현 책임 대상 보고 준비, 계측보고서 자동화가 주요 업무 흐름이다."
    )

    doc.add_heading("날짜별 핵심 업무", level=1)
    timeline = doc.add_table(rows=1, cols=4)
    timeline.style = "Table Grid"
    set_table_widths(timeline, [1.05, 2.0, 1.0, 2.45])
    headers = ["날짜", "대표 업무", "상태", "한 줄 결과"]
    for i, header_text in enumerate(headers):
        cell = timeline.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, LIGHT)
        run = cell.paragraphs[0].add_run(header_text)
        set_font(run, size=9.2, bold=True, color=NAVY)
    for day, text in docs:
        rows = parse_calendar(text)
        if not rows:
            continue
        for idx, row in enumerate(rows):
            cells = timeline.add_row().cells
            values = [day if idx == 0 else "", row[1], row[2], row[3]]
            for c_idx, value in enumerate(values):
                cells[c_idx].text = ""
                run = cells[c_idx].paragraphs[0].add_run(value)
                set_font(run, size=8.8, bold=(c_idx == 0))
                if "완료" == value:
                    set_cell_shading(cells[c_idx], "E6F4EA")
                elif "확인 필요" == value:
                    set_cell_shading(cells[c_idx], "FFF3CD")
    set_repeat_table_header(timeline.rows[0])

    doc.add_heading("상세 업무기록", level=1)
    lead = doc.add_paragraph()
    lead.add_run("아래 기록은 날짜별로 구분하고, 같은 날짜 안에서는 대화별 목적과 결과를 분리해 정리했다.")

    for day, text in docs:
        doc.add_page_break()
        doc.add_heading(day, level=1)
        for kind, value in parse_markdown(text):
            if kind == "h1":
                continue
            if kind == "h2":
                doc.add_heading(value, level=2)
            elif kind == "h3":
                doc.add_heading(value, level=3)
            elif kind == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.1
                run = p.add_run(value)
                set_font(run, size=10)
            elif kind == "p":
                p = doc.add_paragraph()
                run = p.add_run(value)
                set_font(run, size=10.2)
            elif kind == "table":
                add_table_from_markdown(doc, value)

    doc.core_properties.title = "완속충전기 통합 업무기록"
    doc.core_properties.subject = "ChatGPT Work 및 Codex 대화 기반 통합 업무일지"
    doc.core_properties.author = "길앤에스"
    doc.core_properties.keywords = "완속충전기, 업무일지, ChatGPT Work, Codex, MapWebApp, Total_list"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
