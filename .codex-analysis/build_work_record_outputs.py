from __future__ import annotations

from copy import deepcopy
from datetime import date
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\USER\Documents\완속충전기\통합업무기록")
SOURCE = ROOT / "완속충전기_통합업무기록_2026-07-06_현재.docx"
DOCX_OUT = ROOT / "완속충전기_통합업무기록_2026-07-06_현재_업무캘린더교체.docx"
HTML_OUT = ROOT / "완속충전기_통합업무기록_2026-07-06_현재.html"

EVENTS = {
    6: [("상태·마커 동기화", "green")],
    7: [("MapWebApp 3.0 개선", "green")],
    8: [("목록·우선순위 정비", "green")],
    10: [("MapWebApp 4.0", "blue"), ("MXO 얼라이먼트 오류", "orange")],
    13: [("station_no 정합성 기준", "green")],
    14: [("MXO Alignment", "orange"), ("RT-ZHD16 Zero", "orange")],
    15: [("보고서 목록 정비", "green"), ("LMG600 mV/kHz", "orange"), ("LMG600 250V", "orange")],
    20: [("후보·번호·지도 정비", "green")],
    21: [("목록·보고서 검증", "blue"), ("측정 지그 방수", "orange")],
    22: [("월간 보고·To-do", "orange"), ("장비 연결 구성 확인", "orange")],
    23: [("MapWebApp 5.0 개선 공유", "orange")],
    24: [("운전자 연결·촬영 연동", "green"), ("보고서 자동화 구상", "blue")],
}

EVENT_LINKS = {
    (10, "MXO 얼라이먼트 오류"): ("issues", "issue-01"),
    (14, "MXO Alignment"): ("issues", "issue-01"),
    (14, "RT-ZHD16 Zero"): ("issues", "issue-02"),
    (15, "LMG600 mV/kHz"): ("issues", "issue-03"),
    (15, "LMG600 250V"): ("issues", "issue-04"),
    (21, "측정 지그 방수"): ("issues", "issue-05"),
    (22, "장비 연결 구성 확인"): ("issues", "issue-06"),
    (22, "월간 보고·To-do"): ("history", "history"),
}


def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), hex_color)


def set_cell_border(cell, color: str = "D9E2E6", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_calendar_cell(cell, day: int | None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    if day is None:
        shade(cell, "F6F8F9")
        return
    events = EVENTS.get(day, [])
    primary_kind = events[0][1] if events else ""
    shade(cell, {"green": "E7F4EE", "blue": "E9F1FA", "purple": "F0EBF8", "orange": "FFF1DF"}.get(primary_kind, "FFFFFF") if events else "FFFFFF")
    r = p.add_run(str(day))
    set_font(r, 9.5, bold=True, color="174C4A" if events else "5B6870")
    for label, _kind in events:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        r2 = p2.add_run("• " + label)
        set_font(r2, 7.2, bold=True, color="25333A")


def make_docx() -> None:
    doc = Document(SOURCE)
    old_table = doc.tables[1]
    calendar = doc.add_table(rows=6, cols=7)
    calendar.autofit = False
    cal_width = (doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / 7
    for row_idx, row in enumerate(calendar.rows):
        row.height = Inches(0.33 if row_idx == 0 else 0.68)
        for cell in row.cells:
            cell.width = cal_width
    weekdays = ["일", "월", "화", "수", "목", "금", "토"]
    for idx, name in enumerate(weekdays):
        cell = calendar.cell(0, idx)
        shade(cell, "174C4A")
        set_cell_border(cell, "174C4A")
        set_cell_margins(cell, top=50, bottom=50)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        set_font(r, 9, bold=True, color="FFFFFF")
    first_weekday = (date(2026, 7, 1).weekday() + 1) % 7  # Sunday = 0
    for slot in range(35):
        day = slot - first_weekday + 1
        format_calendar_cell(calendar.cell(slot // 7 + 1, slot % 7), day if 1 <= day <= 31 else None)
    old_table._element.addprevious(calendar._tbl)
    old_table._element.getparent().remove(old_table._element)
    doc.save(DOCX_OUT)


def table_html(table, cls: str = "record-table") -> str:
    rows = []
    for r_idx, row in enumerate(table.rows):
        tag = "th" if r_idx == 0 else "td"
        cells = "".join(f"<{tag}>{escape(cell.text).replace(chr(10), '<br>')}</{tag}>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    return f'<div class="table-wrap"><table class="{cls}">{"".join(rows)}</table></div>'


def make_calendar_html() -> str:
    weekdays = ["일", "월", "화", "수", "목", "금", "토"]
    header = "".join(f"<div class='weekday'>{x}</div>" for x in weekdays)
    first_weekday = (date(2026, 7, 1).weekday() + 1) % 7
    cells = []
    for slot in range(35):
        day = slot - first_weekday + 1
        if not 1 <= day <= 31:
            cells.append("<div class='day empty'></div>")
        else:
            events = EVENTS.get(day, [])
            kind = events[0][1] if events else ""
            label_buttons = []
            for label, event_kind in events:
                tab_target, scroll_target = EVENT_LINKS.get(
                    (day, label), ("daily", f"daily-2026-07-{day:02d}")
                )
                label_buttons.append(
                    f"<button class='event {event_kind}' type='button' data-tab-target='{tab_target}' data-scroll-target='{scroll_target}' aria-label='{day}일 상세 내용 보기'>{escape(label).replace('MapWebApp ', 'MapWebApp&nbsp;')}<b aria-hidden='true'>↗</b></button>"
                )
            labels = "".join(label_buttons)
            cells.append(f"<div class='day {kind}'><span class='date'>{day}</span>{labels}</div>")
    legend = "<div class='legend'><span class='legend-green'>● 데이터·자료 정비</span><span class='legend-orange'>● 보고·기술 개선</span><span class='legend-blue'>● 업무 관리·자동화</span></div>"
    return f"<div class='calendar'><div class='calendar-head'>{header}</div><div class='calendar-grid'>{''.join(cells)}</div>{legend}</div>"


def make_change_history_html() -> str:
    return """
<div class='history'>
  <article class='history-card'>
    <div class='version'>MapWebApp 3.0 <span>7/7</span></div>
    <h3>현장 상태와 계측 불가 사유 표시를 개선</h3>
    <ul>
      <li>3kW급, 11/14kW급, 방문 완료 상태의 색상이 서로 비슷해 현장에서 구분하기 어려운 문제를 해소.</li>
      <li>지도 마커와 범례의 색상 체계를 분리해 충전용량과 방문 상태를 바로 읽을 수 있도록 변경.</li>
      <li>계측 불가 사유를 하나만 남길 수 있던 구조를 복수 선택 방식으로 변경.</li>
      <li>출입 불가, 선 없음, 고장, kW 변경, 불법주차, 충전 오류 등 복합 사유를 동시에 기록할 수 있도록 정비.</li>
      <li>선택된 계측 불가 사유와 방문 완료 행의 표시 색상을 맞춰, 지도·현황판·저장 상태가 같은 결과를 보도록 개선.</li>
    </ul>
  </article>
  <article class='history-card'>
    <div class='version'>MapWebApp 4.0 <span>7/10</span></div>
    <h3>최신 목록을 기준으로 지도 데이터 구조를 정비</h3>
    <ul>
      <li>MapWebApp의 충전소 데이터가 최신 <code>Total_list</code>의 번호·우선순위 정보와 같은 기준을 사용하도록 정리.</li>
      <li>앱 기능의 고유 식별자는 <code>id</code>로 통일하고, <code>candidate_id</code>는 지도·검색·필터·CSV 매칭 기준에서 제외.</li>
      <li>방문기록 CSV는 일련번호와 <code>id</code>를 매칭 기준으로 사용하도록 고정해 동일 충전소를 안정적으로 찾게 함.</li>
      <li>지역별로 번호를 수시로 다시 매기지 않고, 전체 현장 출장 종료 뒤 최종 리넘버링을 적용하는 기준을 정리.</li>
      <li>신버전 목록·지도·측정자료가 서로 다른 번호 체계를 갖지 않도록 이후 비교·검증의 기준축을 <code>station_no</code>로 설정.</li>
    </ul>
  </article>
  <article class='history-card'>
    <div class='version'>Total_list · 일련번호 <span>7/7–7/21</span></div>
    <h3>목록을 기준축으로 데이터 연결을 정비</h3>
    <ul>
      <li>4~12지역 충전소의 추가·제거·우선순위 변경 사항을 최신 <code>Total_list</code>에 반영할 기준을 정리.</li>
      <li><code>priority_label</code>의 1순위·2순위를 <code>preset_rank</code>의 <code>rank1</code>·<code>rank2</code>로 일치시킴.</li>
      <li>지역 번호별 1순위는 <code>001</code>부터, 2순위는 <code>201</code>부터 <code>station_no</code>를 다시 부여하는 규칙을 적용.</li>
      <li>새 <code>station_no</code>의 뒤 세 자리를 이용해 <code>station_seq</code>, <code>preset_order</code>, <code>priority_order</code>를 다시 산정.</li>
      <li>변경된 <code>station_no</code>를 MapWebApp의 일련번호와 측정결과보고서의 충전소 번호에 반영하고, 중복·누락·주소·좌표 불일치를 함께 검증.</li>
    </ul>
  </article>
  <article class='history-card'>
    <div class='version'>MapWebApp 5.0 <span>7/23</span></div>
    <h3>현장 운영과 보고용 기능을 통합</h3>
    <ul>
      <li>1순위 후보의 계측이 불가능할 때 바로 대체 방문할 수 있도록 1순위·2순위 운영 체계를 지도에 반영.</li>
      <li>당진·서산부터 목포·해남까지 2순위 후보를 1,692개소에서 2,457개소로 확대해 현장 대체 후보 부족을 보완.</li>
      <li>계측 불가 사유 9개 항목을 복수 기록하고, 방문 완료·모두 사용 중·계측 불가 상태를 색상으로 구분.</li>
      <li>지도 확대 시 일련번호를 표시하고 동일 일련번호의 중복 마커 생성을 막아 중복 집계 위험을 줄임.</li>
      <li>일련번호, 우선순위, 충전용량, 방문 여부, 계측 상태를 조합하는 검색·복합 필터와 지역별 진행 현황 표시를 개선.</li>
    </ul>
  </article>
</div>"""


def make_issue_report_html() -> str:
    issues = [
        ("01", "MXO 5 얼라이먼트 실행 오류", "07.10", "07.14", "채널(C1~C8)에 프로브가 꽂힌 상태에서 Alignment를 실행하면 Id: 1408 오류가 발생.", "모든 채널의 프로브를 제거한 뒤 Alignment를 다시 실행하여 통과."),
        ("02", "전압 프로브(RT-ZHD16) 영점 오차", "07.14", "07.14", "전압 파형이 0V 기준선에 정렬되지 않아 측정 기준점이 흔들림.", "신호 핀과 접지 집게를 직접 쇼트한 뒤 AutoZero를 수행하여 영점 보정."),
        ("03", "LMG600 단위 표시 오류", "07.15", "07.15", "노이즈를 주파수로 인식하거나 전압 범위를 너무 낮게 잡아 mV·kHz로 표시.", "Bandwidth를 Fund. (f1)로, Range를 250V로 고정해 표시 기준을 정상화."),
        ("04", "LMG600 전압 범위 선택 제한", "07.15", "07.15", "Jack 설정이 Usensor여서 안전 제한으로 4V 이상 범위를 선택할 수 없었음.", "노란색 잭 전용 U* 모드로 변경해 250V 범위를 활성화."),
        ("05", "측정 지그 인출부 우천 방수 우려", "07.21", "07.21", "CP/PP 신호선 인출 틈새로 빗물이 들어갈 가능성이 확인됨.", "자기 융착 테이프로 인출부를 밀봉하고 워터 루프를 만들어 빗물 유입 경로를 차단."),
        ("06", "장비별 연결 프로브·클램프 식별 혼선", "07.22", "07.22", "오실로스코프와 파워 어널라이저에 연결하는 프로브·클램프의 수량과 용도가 혼동됨.", "오실로스코프 1개(RT-ZHD16), 파워 어널라이저 2개(전압 리드선·전류 클램프)로 결선 구성을 재확인."),
    ]
    cards = []
    for no, title, occurred, acted, problem, action in issues:
        cards.append(f"""<article class='issue-card' id='issue-{no}'>
  <button class='back-calendar' type='button'>← 달력으로 돌아가기</button>
  <div class='issue-top'><span class='issue-no'>ISSUE {no}</span><span class='issue-dates'><b>발생</b> {occurred}<i>→</i><b>개선</b> {acted}</span><span class='done'>완료</span></div>
  <h3>{escape(title)}</h3>
  <div class='issue-row problem-row'><span>문제</span><p>{escape(problem)}</p></div>
  <div class='issue-row action-row'><span>개선</span><p>{escape(action)}</p></div>
</article>""")
    return "<div class='issues'>" + "".join(cards) + "</div>"


DAILY_IMPROVEMENTS = [
    ("2026-07-06", "상태 저장·마커·현황판 동기화", [
        "문제: 방문 완료·모두 사용 중·계측 불가 상태가 새로고침 후 유지되지 않아 현장 기록이 달라질 가능성 존재.",
        "조치: 상태 저장·복원 흐름 구성 및 재접속 시 마지막 현장 상태 호출.",
        "문제: 지도 마커 색상, 일련번호 라벨, 방문 현황 집계, 저장 데이터가 서로 다른 상태를 표시할 가능성 존재.",
        "조치: 상태 변경 시 지도·라벨·현황 집계·저장 데이터의 상태값 동기화 및 방문 상태 표기 통일.",
        "결과: 새로고침·브라우저 재실행 후 현장 방문·계측 상태 연속 확인.",
    ]),
    ("2026-07-07", "MapWebApp 3.0 표시·사유 입력 개선", [
        "문제: 3kW급·11/14kW급·방문 완료 마커 간 색상 유사로 인한 지도 상태 구분 지연.",
        "조치: 충전용량과 방문 완료 상태의 마커·범례 색상 체계 분리.",
        "문제: 단일 선택 방식으로 인한 복합 계측 불가 사유 기록 한계.",
        "조치: 출입 불가·선 없음·고장·kW 변경·불법주차·충전 오류 등 복수 사유 선택 구조 적용.",
        "결과: 상태 표시 및 계측 불가 사유의 지도·현황판 반영 기준 명확화.",
    ]),
    ("2026-07-08", "4~12지역 목록·우선순위·일련번호 재정비", [
        "문제: group_no 4~12 지역의 추가·제거·우선순위 변경 반영 기준 부재.",
        "조치: priority_label 기준 preset_rank를 rank1·rank2로 통일.",
        "조치: 지역별 station_no 재부여: 1순위 001~, 2순위 201~.",
        "조치: station_no 뒤 세 자리 기준 station_seq·preset_order 재생성 및 priority_order 기준 정비.",
        "결과: Total_list·MapWebApp·측정결과보고서 간 우선순위·일련번호 기준 통일 및 검증 기반 확보.",
    ]),
    ("2026-07-10", "MapWebApp 4.0 데이터 기준 정리", [
        "문제: 목록·지도·방문기록 간 식별자·번호 기준 불일치에 따른 충전소 오매칭 위험.",
        "조치: 앱 고유 식별자를 id로 고정하고 candidate_id를 화면·검색·필터·CSV 처리 대상에서 제외.",
        "조치: 방문기록 CSV의 일련번호→id 연결 기준 정비.",
        "조치: 신버전 Total_list station_no를 지도·측정자료 번호 정합성 검증 기준으로 지정.",
        "결과: MapWebApp 4.0 목록·방문기록·지도 데이터 비교 기준 단일화.",
    ]),
    ("2026-07-13", "station_no·주소·일련번호 정합성 검증", [
        "문제: 목록·지도·측정결과보고서 간 번호·주소 불일치에 따른 결과 통합 오류 위험.",
        "조치: station_no·주소·일련번호 통합 비교 기준 수립.",
        "조치: 번호뿐 아니라 주소·좌표까지 비교하는 동일 충전소 판정 기준 적용.",
        "조치: 중복·누락·자동생성 번호 차이·매칭 불가 항목의 별도 확인 체계 구성.",
        "결과: Total_list 기준 MapWebApp·측정결과보고서 번호 대조 체계 확보.",
    ]),
    ("2026-07-15", "측정결과보고서 목록 정비", [
        "문제: 측정결과보고서 목록의 충전소 표기·정렬 기준 재점검 필요.",
        "조치: 보고서 목록의 충전소 표기와 관리 항목 정비.",
        "결과: 보고서 목록 확인·관리 기준 정리.",
    ]),
    ("2026-07-20", "최신 Total_list를 지도·보고서에 동기화", [
        "문제: 4~12지역 2순위 후보 보강 후 최신 목록과 기존 지도·측정보고서 간 충전소 수·번호 불일치.",
        "조치: 수정 Total_list 기준 측정결과보고서 충전소 번호 재매칭·갱신.",
        "조치: MapWebApp 일련번호를 최신 Total_list station_no와 정합화하고 누락·중복 동시 점검.",
        "조치: 기존 지도 2,633개소와 최신 목록 3,827개소 비교 및 신규 1,194개소 추가 구조 확인.",
        "결과: 후보 추가·제거·우선순위 변경의 목록·지도·보고서 번호 체계 연계.",
    ]),
    ("2026-07-21", "목록·보고서·방문기록 교차 검증", [
        "문제: 파일별 중복·누락·주소·좌표·일련번호 불일치에 따른 진행 현황·계측 결과 집계 오류 위험.",
        "조치: Total_list·측정결과보고서의 주소·위도·경도·충전소 번호·자동생성 번호 교차 비교.",
        "조치: 방문기록 CSV에 목록버전·일련번호·방문 열을 구성하고 MapWebApp 가져오기 형식과 정합화.",
        "조치: 빈·중복 일련번호, 행 수, 헤더 구성의 검증 범위 포함.",
        "결과: 최신 목록·보고서·지도 방문기록의 동일 충전소 지칭 여부 재검증 기반 확보.",
    ]),
    ("2026-07-22", "지도 일련번호 상시 표시 복구", [
        "문제: 마커 선택 후 팝업 종료 시 확대 상태의 일련번호 라벨 소실 및 지도 이동 전 미복구.",
        "문제: 내부 상태값과 실제 Leaflet 툴팁 열림 상태 간 불일치.",
        "조치: 현재 필터·레이어 포함 여부·확대 수준·지도 범위 기준의 라벨 표시 판단 로직 적용.",
        "조치: 실제 툴팁 존재·열림 상태 검사 및 누락 라벨 재표시 로직 적용.",
        "결과: 마커 선택 해제 후 지도상 일련번호의 지속 표시.",
    ]),
    ("2026-07-23", "MapWebApp 5.0 현장 운영 기능 통합", [
        "문제: 1순위 충전소 계측 불가 시 지역별 목표 수량 충족을 위한 대체 후보 부족.",
        "조치: 1순위·2순위 운영 체계 도입 및 2순위 후보 1,692개소→2,457개소 확대.",
        "조치: 계측 불가 사유 9개 항목 복수 기록 및 방문·계측 상태 색상 구분.",
        "조치: 확대 시 일련번호 표시, 중복 마커 방지, 번호·우선순위·용량·방문·계측 상태 복합 필터 적용.",
        "결과: 대체 후보 탐색·진행 현황·데이터 품질의 현장 지도 통합 관리.",
    ]),
    ("2026-07-24", "운전자 연결·촬영 연동 및 보고서 자동화 구상", [
        "문제: 일일 계측 CSV를 측정결과보고서에 반복 반영하기 위한 입력 기준·자동화 방식 미확정.",
        "조치: 퍼센트 변환·제조년월 입력을 포함한 계측보고서 자동 작성 방식 검토.",
        "조치: 운전자 연결·제원 사진 촬영 화면과 Google Apps Script 구성 파일 생성.",
        "결과: 운전자 연결·촬영 연동 및 보고서 자동화 관련 파일 생성 확인. 배포 주소·권한 설정·실제 통신 시험은 추가 확인 필요.",
    ]),
]


def make_daily_history_html() -> str:
    blocks = []
    for day, title, items in DAILY_IMPROVEMENTS:
        problem_no = 0
        action_no = 0
        lines = []
        for item in items:
            kind = "other"
            prefix, content = (item.split(":", 1) + [""])[:2] if ":" in item else ("내용", item)
            prefix = prefix.strip()
            content = content.strip()
            if prefix in ("문제", "원인"):
                problem_no += 1
                label = f"문제 {problem_no}"
                kind = "problem"
            elif prefix in ("조치", "반영"):
                action_no += 1
                label = f"개선 {action_no}"
                kind = "action"
            elif prefix == "결과":
                label = "결과"
                kind = "result"
            else:
                label = prefix
            lines.append(f"<li class='{kind}'><span class='step-label'>{escape(label)}</span><span>{escape(content)}</span></li>")
        blocks.append(f"<details class='daily' id='daily-{day}' open><summary><span>{day}</span><strong>{escape(title)}</strong></summary><div class='improvement'><button class='back-calendar' type='button'>← 달력으로 돌아가기</button><ul class='detail-steps'>{''.join(lines)}</ul></div></details>")
    return "".join(blocks)


def make_html() -> None:
    sections = make_daily_history_html()
    html = f"""<!doctype html>
<html lang='ko'>
<head>
<meta charset='utf-8'>
<meta name='viewport' content='width=device-width, initial-scale=1'>
<title>완속충전기 통합 업무기록</title>
<style>
:root{{--bg:#171717;--card:#202020;--line:#404040;--text:#f4f4f4;--muted:#b4b4b4;--green:#72d99c;--orange:#ff9e55;--blue:#acd7ff}}
*{{box-sizing:border-box}} body{{margin:0;background:var(--bg);color:var(--text);font-family:'Malgun Gothic','Apple SD Gothic Neo',Arial,sans-serif;line-height:1.6}}
.page{{width:100%;max-width:1440px;margin:0 auto;padding:0 42px 56px}} .hero{{display:flex;align-items:baseline;justify-content:space-between;gap:24px;padding:0 0 23px;border-bottom:1px solid var(--line)}}
.eyebrow{{display:none}} h1{{flex:1;min-width:0;white-space:nowrap;font-size:30px;line-height:1.2;margin:0;font-weight:800;letter-spacing:-.04em}} .hero p{{flex:0 0 auto;margin:0;color:var(--muted);font-size:16px;font-weight:600}}
.panel{{margin-top:22px}} h2{{margin:0 0 18px;font-size:21px;letter-spacing:-.02em}}
.table-wrap{{overflow-x:auto}} table{{border-collapse:collapse;width:100%;font-size:14px}} th{{background:#292929;color:#fff;font-weight:700}} th,td{{border:1px solid var(--line);padding:10px 12px;vertical-align:top;text-align:left}} .metadata th{{width:22%;white-space:nowrap}} .metadata td{{font-weight:600}}
.calendar{{border:0;border-radius:0;overflow:visible}} .calendar-head,.calendar-grid{{display:flex;flex-wrap:wrap;gap:9px}} .calendar-head{{margin-bottom:9px}} .weekday{{flex:0 1 calc((100% - 54px) / 7);min-width:0;color:var(--muted);text-align:center;font-weight:700;padding:0;font-size:17px}} .day{{box-sizing:border-box;flex:0 1 calc((100% - 54px) / 7);min-width:0;min-height:132px;padding:13px 10px;border:1px solid var(--line);border-radius:13px;background:var(--card);display:flex;flex-direction:column;gap:7px}} .day.empty{{border-color:transparent;background:transparent}} .date{{font-weight:800;color:#fff;font-size:21px;line-height:1}} .event{{display:block;width:100%;padding:3px 4px;border:1px solid transparent;border-radius:6px;background:transparent;color:#f7f7f7;cursor:pointer;white-space:nowrap;font-family:inherit;text-align:left;font-size:clamp(12px,.9vw,14px);font-weight:800;line-height:1.35;letter-spacing:-.055em;transition:background .16s ease,border-color .16s ease,transform .16s ease}} .event:hover,.event:focus-visible{{background:#2b3330;border-color:#4b8f67;outline:none;transform:translateX(2px)}} .event b{{float:right;margin-left:3px;color:#8fe6b2;font-size:11px;opacity:.8}} .event::before{{content:'';display:inline-block;width:10px;height:10px;border-radius:50%;background:var(--green);margin:0 6px 1px 0}} .event.orange::before{{background:var(--orange)}} .event.blue::before{{background:var(--blue)}} .event.purple::before{{background:var(--orange)}} .daily.jump-target{{outline:2px solid var(--green);outline-offset:3px;box-shadow:0 0 0 7px rgba(112,225,160,.13)}}
.legend{{display:flex;flex-wrap:wrap;gap:18px;margin-top:20px;color:#d1d1d1;font-size:14px;font-weight:700}} .legend-green{{color:#d1d1d1}} .legend-orange{{color:#d1d1d1}} .legend-blue{{color:#d1d1d1}} .legend-green::first-letter{{color:var(--green)}} .legend-orange::first-letter{{color:var(--orange)}} .legend-blue::first-letter{{color:var(--blue)}} .foot{{color:var(--muted);font-size:13px;margin-top:26px}}
.history{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:14px}} .history-card{{position:relative;border:1px solid var(--line);border-radius:13px;background:linear-gradient(135deg,#222 0%,#1d1d1d 100%);padding:19px 21px 20px;overflow:hidden}} .history-card::before{{content:'';position:absolute;inset:0 auto 0 0;width:4px;background:var(--blue)}} .history-card:nth-child(odd)::before{{background:var(--green)}} .history-card h3{{margin:12px 0 14px;font-size:19px;line-height:1.3;letter-spacing:-.04em}} .version{{display:flex;align-items:center;justify-content:space-between;color:var(--green);font-size:12px;font-weight:900;letter-spacing:.04em}} .version span{{color:#aeb6b9;font-weight:700;letter-spacing:0}} .history-card ul{{margin:0;padding:0;list-style:none;color:#e2e2e2;font-size:13px;line-height:1.62}} .history-card li{{position:relative;padding:8px 10px 8px 24px;border-radius:7px;background:#202b27}} .history-card li+li{{margin-top:6px}} .history-card li::before{{content:'✓';position:absolute;left:9px;color:var(--green);font-weight:900}} .history-card:nth-child(even) li{{background:#20262c}} .history-card:nth-child(even) li::before{{color:var(--blue)}} .history-card code{{font-family:Consolas,'Malgun Gothic',monospace;font-size:.92em;color:#baddff}}
.issues{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:14px}} .issue-card{{position:relative;border:1px solid var(--line);border-radius:13px;background:linear-gradient(135deg,#222 0%,#1d1d1d 100%);padding:19px 21px 20px;overflow:hidden}} .issue-card::before{{content:'';position:absolute;inset:0 auto 0 0;width:4px;background:var(--green)}} .issue-top{{display:flex;align-items:center;gap:10px}} .issue-no{{color:var(--green);font-size:12px;font-weight:900;letter-spacing:.04em}} .issue-dates{{margin-left:auto;display:flex;align-items:center;gap:6px;color:#aeb6b9;font-size:12px;font-weight:700}} .issue-dates b{{color:#e4e4e4;font-size:11px}} .issue-dates i{{color:#6e777a;font-style:normal}} .done{{border:1px solid #347955;border-radius:20px;padding:3px 9px;color:var(--green);font-size:12px;font-weight:800;white-space:nowrap}} .issue-card h3{{margin:12px 0 15px;font-size:19px;line-height:1.28;letter-spacing:-.04em}} .issue-row{{display:grid;grid-template-columns:50px minmax(0,1fr);gap:10px;padding:10px 11px;border-radius:8px}} .issue-row+ .issue-row{{margin-top:7px}} .issue-row>span{{font-size:12px;font-weight:900;line-height:1.65}} .issue-row p{{margin:0;color:#e7e7e7;font-size:13px;line-height:1.65}} .problem-row{{background:#2b2722}} .problem-row>span{{color:#ffb06a}} .action-row{{background:#202c27}} .action-row>span{{color:var(--green)}} .back-calendar{{display:none;margin:0 0 12px;padding:6px 10px;border:1px solid #4b8f67;border-radius:7px;background:#202c27;color:#8fe6b2;cursor:pointer;font:inherit;font-size:13px;font-weight:800}} .from-calendar .back-calendar{{display:inline-flex;align-items:center;gap:5px}} .back-calendar:hover,.back-calendar:focus-visible{{background:#2b4236;outline:none}}
.tabbed-report{{margin-top:42px;border-top:1px solid var(--line);padding-top:20px}} .tabs{{display:grid;grid-template-columns:repeat(3,1fr);border-bottom:1px solid var(--line)}} .tab{{appearance:none;border:0;border-bottom:3px solid transparent;background:transparent;color:var(--muted);cursor:pointer;padding:13px 12px 12px;font:inherit;font-size:16px;font-weight:800;letter-spacing:-.03em;text-align:center}} .tab:hover{{color:#fff}} .tab.active{{border-bottom-color:var(--green);color:#fff}} .tab-panel{{display:none;padding-top:22px}} .tab-panel.active{{display:block}} .tab-panel h2{{margin-bottom:4px}}
.daily{{position:relative;border:1px solid var(--line);border-radius:13px;background:linear-gradient(135deg,#222 0%,#1d1d1d 100%);margin-top:12px;overflow:hidden}} .daily::before{{content:'';position:absolute;inset:0 auto 0 0;width:4px;background:var(--blue)}} .daily summary{{display:flex;align-items:center;gap:13px;cursor:pointer;padding:16px 20px;font-size:16px;font-weight:700;list-style:none;background:#202020}} .daily summary::-webkit-details-marker{{display:none}} .daily summary span{{display:inline-flex;align-items:center;justify-content:center;min-width:98px;border:1px solid #347955;border-radius:18px;padding:4px 10px;color:var(--green);font-size:12px;font-weight:900;letter-spacing:.02em}} .daily summary strong{{font-size:17px;letter-spacing:-.035em}} .daily summary::after{{content:'+';margin-left:auto;color:var(--muted);font-size:21px;line-height:18px}} .daily[open] summary{{border-bottom:1px solid var(--line)}} .daily[open] summary::after{{content:'−'}} .improvement{{padding:15px 20px 19px}} .detail-steps{{margin:0;padding:0;list-style:none;color:#e7e7e7;font-size:14px;line-height:1.62}} .detail-steps li{{display:flex;gap:12px;align-items:flex-start;padding:10px 11px;border-radius:8px;background:#252525}} .detail-steps li+li{{margin-top:7px}} .detail-steps li.problem{{background:#2b2722}} .detail-steps li.action{{background:#202c27}} .detail-steps li.result{{background:#20262c}} .step-label{{flex:0 0 64px;font-size:12px;font-weight:900;line-height:1.9}} .detail-steps li.problem .step-label{{color:#ffb06a}} .detail-steps li.action .step-label{{color:var(--green)}} .detail-steps li.result .step-label{{color:var(--blue)}}
.tabbed-report{{padding-top:24px}} .tab{{padding:16px 12px 15px;font-size:19px;font-weight:900}} .tab-panel{{padding-top:26px}}
.issues,.history{{gap:16px}} .issue-card,.history-card{{padding:22px 24px 24px}} .issue-no,.version{{font-size:13px}} .issue-dates{{gap:7px;font-size:13px}} .issue-dates b{{font-size:12px}} .done{{padding:4px 10px;font-size:13px}} .issue-card h3,.history-card h3{{margin-top:14px;margin-bottom:18px;font-size:22px;line-height:1.32}} .issue-row{{grid-template-columns:56px minmax(0,1fr);gap:12px;padding:12px 13px}} .issue-row+ .issue-row{{margin-top:9px}} .issue-row>span{{font-size:13px;line-height:1.75}} .issue-row p{{font-size:15px;line-height:1.75}}
.history-card ul{{font-size:15px;line-height:1.72}} .history-card li{{padding:10px 12px 10px 28px}} .history-card li+li{{margin-top:8px}} .history-card li::before{{left:11px}}
.daily{{margin-top:14px}} .daily summary{{gap:15px;padding:19px 23px}} .daily summary span{{min-width:108px;padding:5px 11px;font-size:13px}} .daily summary strong{{font-size:20px}} .daily summary::after{{font-size:24px}} .improvement{{padding:18px 23px 22px}} .detail-steps{{font-size:15px;line-height:1.72}} .detail-steps li{{gap:14px;padding:12px 14px}} .detail-steps li+li{{margin-top:9px}} .step-label{{flex-basis:70px;font-size:13px;line-height:1.95}}
.tab{{font-size:21px;letter-spacing:-.025em}} .issue-card h3,.history-card h3{{font-size:24px}} .issue-row p,.history-card ul{{font-size:17px;line-height:1.78}} .issue-row>span{{font-size:19px;line-height:1.8}} .issue-row{{grid-template-columns:94px minmax(0,1fr)}} .daily summary strong{{font-size:22px}} .detail-steps{{font-size:17px;line-height:1.78}} .detail-steps li{{padding:15px 18px}} .step-label{{flex:0 0 94px;font-size:19px;font-weight:900;line-height:1.8;letter-spacing:-.04em}}
@media(max-width:760px){{.page{{padding:0 14px 36px}}.hero{{padding-bottom:18px;align-items:flex-start;gap:8px;flex-direction:column}}h1{{font-size:25px}}.hero p{{font-size:14px}}.calendar-head,.calendar-grid{{gap:5px}}.weekday{{font-size:12px}}.day{{min-height:82px;padding:8px 7px;border-radius:9px;gap:7px}}.date{{font-size:15px}}.event{{font-size:10px}}.event::before{{width:7px;height:7px;margin-right:4px}}.history,.issues{{grid-template-columns:1fr}}.issue-top{{flex-wrap:wrap}}.issue-dates{{margin-left:0;order:3;width:100%}}.tab{{padding:12px 5px;font-size:13px}}.daily summary{{padding:15px 16px;gap:9px}}.daily summary span{{min-width:90px;padding:4px 7px;font-size:11px}}.daily summary strong{{font-size:16px}}.detail-steps{{font-size:14px}}.detail-steps li{{padding:10px}}.calendar::after{{font-size:11px;line-height:1.8}}table{{font-size:12px}}th,td{{padding:8px}}}}
</style>
</head>
<body><main class='page'>
<header class='hero'><div class='eyebrow'>PROJECT WORK RECORD</div><h1>1st Monthly Report</h1><p>EV Charger Field Inspection · July 2026</p></header>
<section class='panel'>{make_calendar_html()}</section>
<section class='panel tabbed-report'>
  <div class='tabs' role='tablist' aria-label='업무 기록 분류'>
    <button class='tab active' type='button' role='tab' aria-selected='true' data-target='issues'>기술 이슈·조치 리포트</button>
    <button class='tab' type='button' role='tab' aria-selected='false' data-target='history'>주요 개선 이력</button>
    <button class='tab' type='button' role='tab' aria-selected='false' data-target='daily'>날짜별 개선 내용</button>
  </div>
  <div class='tab-panel active' id='issues' role='tabpanel'><p class='foot' style='margin-top:0;margin-bottom:16px'>6월 29일 필드 점검 프로젝트 시작 이후 발생한 주요 장비·물리 연결 이슈와 조치 내역입니다. 현재 6건 모두 해결되어 정상 측정 단계로 전환했습니다.</p>{make_issue_report_html()}</div>
  <div class='tab-panel' id='history' role='tabpanel'>{make_change_history_html()}</div>
  <div class='tab-panel' id='daily' role='tabpanel'>{sections}</div>
</section>
<p class='foot'>본 기록은 기존 통합 업무기록의 날짜별 상세 표를 그대로 옮기고, 핵심 업무 표를 월간 캘린더로 재구성한 버전입니다.</p>
</main><script>
function activateTab(target) {{
  document.querySelectorAll('.tab').forEach((item) => {{
    const active = item.dataset.target === target;
    item.classList.toggle('active', active);
    item.setAttribute('aria-selected', active ? 'true' : 'false');
  }});
  document.querySelectorAll('.tab-panel').forEach((panel) => panel.classList.toggle('active', panel.id === target));
}}
document.querySelectorAll('.tab').forEach((tab) => {{
  tab.addEventListener('click', () => activateTab(tab.dataset.target));
}});
document.querySelectorAll('[data-scroll-target]').forEach((item) => {{
  item.addEventListener('click', () => {{
    const target = document.getElementById(item.dataset.scrollTarget);
    if (!target) return;
    activateTab(item.dataset.tabTarget);
    target.open = true;
    requestAnimationFrame(() => {{
      document.querySelectorAll('.from-calendar').forEach((card) => card.classList.remove('from-calendar'));
      target.classList.add('from-calendar');
      target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
      target.classList.add('jump-target');
      window.setTimeout(() => target.classList.remove('jump-target'), 1600);
    }});
  }});
}});
document.querySelectorAll('.back-calendar').forEach((button) => {{
  button.addEventListener('click', () => {{
    const calendar = document.querySelector('.calendar');
    if (!calendar) return;
    calendar.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
  }});
}});
</script></body></html>"""
    HTML_OUT.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    ROOT.mkdir(exist_ok=True)
    make_docx()
    make_html()
    print(DOCX_OUT)
    print(HTML_OUT)
