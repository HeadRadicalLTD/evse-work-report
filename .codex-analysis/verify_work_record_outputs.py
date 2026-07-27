from pathlib import Path
from docx import Document

root = Path(r"C:\Users\USER\Documents\완속충전기\통합업무기록")
docx_path = root / "완속충전기_통합업무기록_2026-07-06_현재_업무캘린더교체.docx"
html_path = root / "완속충전기_통합업무기록_2026-07-06_현재.html"
doc = Document(docx_path)
print("DOCX exists:", docx_path.exists(), "bytes:", docx_path.stat().st_size)
print("tables:", len(doc.tables), "calendar:", len(doc.tables[1].rows), "x", len(doc.tables[1].columns))
print("calendar header:", [c.text for c in doc.tables[1].rows[0].cells])
print("day 8:", doc.tables[1].cell(2, 3).text.replace("\n", " / "))
text = html_path.read_text(encoding="utf-8")
marks = [
    "상태·마커 동기화", "MapWebApp 3.0 개선", "목록·우선순위 정비", "Map 4.0 · MXO 이슈", "station_no 정합성 기준",
    "RT-ZHD16 AutoZero", "LMG600 250V 설정", "후보·번호·지도 정비", "방수·목록 검증",
    "연결구성·보고 정리", "MapWebApp 5.0 개선 공유", "업무일지·자동화 구상",
]
print("HTML exists:", html_path.exists(), "bytes:", html_path.stat().st_size)
print("calendar marks:", sum(mark in text for mark in marks))
print("detail records:", text.count("<details class='daily' open>"))
print("version history cards:", text.count("class='history-card'"))
print("legacy status header absent:", "<th>상태</th>" not in text)
print("improvement lists:", text.count("<ul class='detail-steps'>"))
print("technical issue cards:", text.count("class='issue-card'"))
print("all six completed:", text.count(">완료</span>") == 6)
print("tab headers:", all(label in text for label in ["기술 이슈·조치 리포트", "주요 개선 이력", "날짜별 개선 내용"]))
print("tab panels:", text.count("class='tab-panel'") == 2 and text.count("class='tab-panel active'") == 1)
print("tab interaction script:", "tab.addEventListener('click'" in text)
print("monthly report title:", "1st Monthly Report" in text)
